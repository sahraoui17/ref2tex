import sys
import getopt
import re
import bibtexparser
from bibtexparser.bibdatabase import BibDatabase
import io 
from docx import Document

# read a DOCX file and split it into two TXT files (References and Text files)
def docx2txt(InputDocx): 
    '''convert a DOCX file into two TXT files to be saved localy'''
    document = Document(InputDocx)
    # get index at which "References" heading starts 
    for index, paragraph in enumerate(document.paragraphs): 
        if paragraph.style.name in ('Heading', 'Heading 1', 'Heading 2')\
                and  ('Reference' or 'reference') in paragraph.text: 
            get_index = index
    
    # Extract lines after and before "References" heaading as lists of lines
    RefLibList = [paragraph.text for paragraph in document.paragraphs[get_index:]]
    TextInputList = [paragraph.text for paragraph in document.paragraphs[:get_index]]
    
    # Converting the lists into text streams
    TextInput_io = io.StringIO()
    RefLib_io = io.StringIO()
    [TextInput_io.write(i) for i in TextInputList]
    [RefLib_io.write(i) for i in RefLibList]
    
    # Saving the two seperated text streams into local .txt files
    with open('InputText.txt','wt', encoding='utf-8',) as InputText: 
        for i in TextInputList: 
            InputText.write(i +'\n')
    with open('InputRef.txt','wt', encoding='utf-8',) as InputRef: 
        for i in RefLibList: 
            InputRef.write(i +'\n')


# given a file that contains a list of IEEE references, return the reference line
def get_reference_line(referencesFile,numberInBrackets):   
    with open(referencesFile, "rt",encoding="utf-8") as inputFile:
        for line in inputFile:
            if line.startswith(numberInBrackets):
                return line
    return 0

def get_title(referenceLine):
    if(len(referenceLine)>0):#skip empty lines
        try:
            title=re.search('“(.*),”', referenceLine).group(1)
            return title
        except IndexError:
            print(referenceLine,'was not processed')
            return 0

# given a reference number, get its title, and return the correspondng bib key
def get_reference_key(referenceLine,bib_database):
    title=get_title(referenceLine) # get the reference title
    title=title.replace('’',"'")  # MS word have different ' with bib/txt.  ’ and ' are different    
    if title[-1]=='.':# some references has . (dot) at the end, it must be removed before searching in bib file
        title=title[:-1]
    try:

        return next(entry for entry in bib_database.entries if entry['title'].replace('{','').replace('}','')==title)['ID'] # A temporary hacky way to remove the curly braces from the Bibtex titles' capitalized words, to avoid the "Reference NOT FOUND in bib file" Error  
    except StopIteration:
        print('Reference NOT FOUND in bib file:',title)
        return False

# read a document that contain inline citation in IEEE, i.e. ([7]) and relace it with tex citation command \cite{key_in_bib_file}
def ieee2tex(docxFile,bibFile,outputTexFile='output_cited.tex'):
    docx2txt(docxFile)
    bib_database = BibDatabase()
    with open(bibFile,encoding="utf-8") as bibtex_file:
        bib_database = bibtexparser.load(bibtex_file)

    with open('InputText.txt', "rt",encoding="utf-8") as inputFile:
        with open(outputTexFile, "wt",encoding="utf-8") as outputFile:
            for line in inputFile:
                citations= re.findall(r'\[\d+\]',line)
                for citation in citations:
                    try:
                        line=line.replace(citation,'\cite{'+get_reference_key(get_reference_line('InputRef.txt',citation),bib_database)+'}')
                    except StopIteration:
                        print('reference not found')    
                outputFile.write(line)


def main(docxFile,bibFile,outputFile='cited_tex.tex'):
    import os.path
   
    # check if inputFile exists
    if not os.path.isfile(docxFile):
        print(docxFile,'does not exist')
        sys.exit()
    # check if bibFile exists 
    if not os.path.isfile(bibFile):
        print(bibFile,'does not exist')
        sys.exit()
    ieee2tex(docxFile,bibFile,outputFile)
    
if __name__ == "__main__":

    for opt, arg in getopt.getopt(sys.argv[1:],'i:b:o:',[])[0]:
        print(opt,arg)
        if opt=='-i':
           docxFile=arg
        if opt=='-b':
           bibFile=arg
        if opt=='-o':
           outputFile=arg
        
    try:
        main(docxFile,bibFile,outputFile)
    except NameError:
        main(docxFile,bibFile)
